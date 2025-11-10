# pages/05_Auto_Article.py
# Auto-Article: builds a full paper draft from app outputs.
# Modes:
#   - Template (offline, no model needed)
#   - Optional Offline LLM via llama.cpp (no API). If a GGUF exists at models/<file>.gguf, it can refine the draft.
#
# Soft deps (optional): llama-cpp-python, python-docx

import os, io, json, textwrap, datetime
import numpy as np
import pandas as pd
import streamlit as st

# ---------- Page ----------
st.set_page_config(
    page_title="📝 Auto-Article (from analyses)",
    page_icon="📝",
    layout="wide"
)
st.title("📝 Auto-Article — Draft a Paper from Your Analyses")

# ---------- Helpers: width handling (new API first, fallback old) ----------
def show_df(df: pd.DataFrame):
    try:
        st.dataframe(df, width="stretch")
    except TypeError:
        st.dataframe(df, use_container_width=True)

# ---------- Try optional local LLM (llama.cpp) ----------
def find_local_gguf():
    """Return a model path if we can find a GGUF under ./models; else None."""
    candidates = []
    mdir = os.path.join(os.getcwd(), "models")
    if os.path.isdir(mdir):
        for fn in os.listdir(mdir):
            if fn.lower().endswith(".gguf"):
                candidates.append(os.path.join(mdir, fn))
    return candidates[0] if candidates else None

def local_llm_available():
    """Check if llama_cpp is importable and a GGUF model exists."""
    try:
        import llama_cpp  # noqa
        return find_local_gguf() is not None
    except Exception:
        return False

def refine_with_local_llm(markdown_seed: str, target_venue: str, tone: str, length: str) -> str:
    """
    Use llama-cpp-python with a local GGUF (no internet, no API) to polish the article draft.
    Keeps things conservative (low temperature). If anything fails, return the seed unchanged.
    """
    try:
        from llama_cpp import Llama
        model_path = find_local_gguf()
        if not model_path:
            return markdown_seed

        # Small context/threads to be gentle on CPU/RAM
        llm = Llama(
            model_path=model_path,
            n_ctx=4096,
            n_threads=max(1, os.cpu_count() // 2),
            verbose=False
        )

        sys_prompt = (
            "You are an expert academic writer. Improve the following Markdown draft into a well-structured, "
            "clear article. Preserve Markdown headings, keep it coherent, tighten prose, and improve transitions. "
            f"Target venue: {target_venue}. Tone: {tone}. Length target: {length}. "
            "Do not invent data or results beyond what is given."
        )
        user_prompt = "Here is the draft to refine:\n\n" + markdown_seed

        # Generic chat template for instruct models (TinyLlama, Llama, Mistral chat variants)
        messages = [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ]
        # llama-cpp expects a single prompt string; we can format our own
        prompt = (
            "### System\n" + sys_prompt + "\n\n"
            "### User\n" + user_prompt + "\n\n"
            "### Assistant\n"
        )

        out = llm(
            prompt=prompt,
            max_tokens=1200,
            temperature=0.2,
            top_p=0.9,
            repeat_penalty=1.1,
            stop=["### User", "### System"]
        )
        text = out.get("choices", [{}])[0].get("text", "").strip()
        return text if text else markdown_seed
    except Exception as e:
        st.warning(f"Local LLM refinement unavailable: {e}")
        return markdown_seed

# ---------- Collect context from session or uploads ----------
st.markdown("#### 1) Inputs to build the article")
left, right = st.columns([1.2, 1])

with left:
    # Pull from earlier pages if they saved into session_state
    df = st.session_state.get("df")
    hark_top = st.session_state.get("harking_top")
    multix_top = st.session_state.get("multix_top")
    reg_summary = st.session_state.get("reg_summary")
    reg_coefs = st.session_state.get("reg_coefs")

    # Fallback uploads (only if session is empty)
    st.markdown("Provide any **missing** pieces below (only if not already in memory).")
    up_df = st.file_uploader("Upload full dataset CSV (optional, for data section)", type=["csv"], key="aa_df")
    if up_df is not None:
        try:
            df = pd.read_csv(up_df)
        except Exception as e:
            st.error(f"Could not read dataset CSV: {e}")

    up_hark = st.file_uploader("Upload HARKing Top-K CSV (optional)", type=["csv"], key="aa_hark")
    if (hark_top is None) and (up_hark is not None):
        try:
            hark_top = pd.read_csv(up_hark)
        except Exception as e:
            st.error(f"Could not read HARKing CSV: {e}")

    up_multix = st.file_uploader("Upload Multivariate HARKing Top-K CSV (optional)", type=["csv"], key="aa_multix")
    if (multix_top is None) and (up_multix is not None):
        try:
            multix_top = pd.read_csv(up_multix)
        except Exception as e:
            st.error(f"Could not read MultiX CSV: {e}")

    up_reg_txt = st.file_uploader("Upload Regression Summary TXT (optional)", type=["txt"], key="aa_regtxt")
    if (reg_summary is None) and (up_reg_txt is not None):
        try:
            reg_summary = up_reg_txt.read().decode("utf-8", errors="ignore")
        except Exception as e:
            st.error(f"Could not read regression summary TXT: {e}")

    up_reg_coef = st.file_uploader("Upload Regression Coefficients CSV (optional)", type=["csv"], key="aa_regcoef")
    if (reg_coefs is None) and (up_reg_coef is not None):
        try:
            reg_coefs = pd.read_csv(up_reg_coef)
        except Exception as e:
            st.error(f"Could not read coefficients CSV: {e}")

with right:
    st.markdown("##### Preview detected inputs")
    if isinstance(df, pd.DataFrame):
        st.caption("Dataset sample")
        show_df(df.head(8))
    if isinstance(hark_top, pd.DataFrame):
        st.caption("HARKing Top-K sample")
        show_df(hark_top.head(8))
    if isinstance(multix_top, pd.DataFrame):
        st.caption("MultiX Top-K sample")
        show_df(multix_top.head(8))
    if isinstance(reg_coefs, pd.DataFrame):
        st.caption("Regression coefficients sample")
        show_df(reg_coefs.head(8))
    if isinstance(reg_summary, str) and reg_summary.strip():
        st.caption("Regression summary present ✓")

# ---------- Article configuration ----------
st.markdown("#### 2) Article configuration")

c1, c2, c3 = st.columns([1.3, 1, 1])
title = c1.text_input("Title", value="Measurement, Correlation, and Confirmation: A Guided Exploration with HARKing & OLS", key="aa_title")
target_venue = c2.selectbox("Target style", ["Academic (econ)", "Academic (applied econ)", "Academic (policy)", "Blog / explainer"], index=0, key="aa_style")
length = c3.selectbox("Length", ["Short (~800 words)", "Standard (~1500–2000 words)", "Long (~3000–4000 words)"], index=1, key="aa_len")

c4, c5 = st.columns([1.5, 1])
sections = c4.multiselect(
    "Include sections",
    ["Abstract", "Introduction", "Data", "Methods", "Results", "Robustness & Diagnostics", "Limitations", "Conclusion", "References"],
    default=["Abstract","Introduction","Data","Methods","Results","Robustness & Diagnostics","Limitations","Conclusion"],
    key="aa_sections"
)
tone = c5.selectbox("Tone", ["Neutral-academic", "Didactic/teaching", "Critical/replication"], index=1, key="aa_tone")

st.markdown("##### Optional metadata")
colA, colB, colC = st.columns([1,1,1])
authors = colA.text_input("Authors", value="Your Name", key="aa_authors")
affil = colB.text_input("Affiliation", value="Your Institution", key="aa_affil")
date_str = colC.text_input("Date", value=datetime.date.today().isoformat(), key="aa_date")

st.markdown("##### Narrative options")
narr1, narr2 = st.columns([1,1])
include_harking_warning = narr1.checkbox("Emphasize multiple-testing/HARKing warnings", True, key="aa_warn")
include_fwll_note = narr2.checkbox("Explain FWL partial regression (main X net of controls)", True, key="aa_fwllexp")

st.markdown("---")

# ---------- Build structured context ----------
def summarize_df(df: pd.DataFrame, max_cols=12):
    parts = []
    try:
        n, k = df.shape
        parts.append(f"Dataset has {n:,} rows and {k:,} columns.")
        num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        dt_cols = [c for c in df.columns if np.issubdtype(df[c].dtype, np.datetime64)]
        obj_cols = [c for c in df.columns if df[c].dtype == "object"]
        if num_cols: parts.append(f"Numeric columns (sample): {', '.join(num_cols[:max_cols])}")
        if dt_cols: parts.append(f"Datetime-like columns: {', '.join(dt_cols[:max_cols])}")
        if obj_cols: parts.append(f"Categoricals/objects (sample): {', '.join(obj_cols[:max_cols])}")
    except Exception:
        pass
    return " ".join(parts) if parts else "Dataset summary unavailable."

def summarize_top_hits(df_top: pd.DataFrame, label: str, n=5):
    if df_top is None or df_top.empty:
        return f"No {label} results available."
    head = df_top.head(n).copy()
    cols = [c for c in head.columns if c.lower() in {"rank","y","x","x_main","tfx","tfy","ty","tx","lag","p","p_main","r2","n","beta_main","b1","controls","filter"}]
    if not cols:
        cols = head.columns[:min(8, len(head.columns))]
    return head[cols].to_markdown(index=False)

def default_methods(include_fwll: bool):
    txt = []
    txt.append("We combine exploratory specification search (HARKing) with confirmatory ordinary least squares (OLS). ")
    txt.append("The HARKing module grid-searches simple models across transforms, lags, and bin-slice filters, ranking by p-value and filtering by minimum R². ")
    txt.append("Confirmatory OLS allows multivariate specifications, diagnostic checks, and multicollinearity screening (VIF). ")
    if include_fwll:
        txt.append("For multivariate models, we use Frisch–Waugh–Lovell partial regression plots to visualize the net association of the main regressor with the outcome after removing linear effects of controls.")
    return "".join(txt)

def default_limitations():
    return (
        "Results from exploratory search are vulnerable to multiple-testing artifacts and researcher degrees of freedom. "
        "OLS relies on linearity, exogeneity, and homoskedasticity assumptions; violations (omitted variables, reverse causality, measurement error) can bias estimates. "
        "Time-series features (trends, unit roots) can inflate significance if not addressed. These outputs are educational and hypothesis-generating, not definitive."
    )

context = {
    "dataset_summary": summarize_df(df) if isinstance(df, pd.DataFrame) else "Dataset not provided.",
    "hark_summary_md": summarize_top_hits(hark_top, "HARKing") if isinstance(hark_top, pd.DataFrame) else "HARKing results not provided.",
    "multix_summary_md": summarize_top_hits(multix_top, "Multivariate HARKing") if isinstance(multix_top, pd.DataFrame) else "Multivariate HARKing results not provided.",
    "reg_summary_text": reg_summary if isinstance(reg_summary, str) else "Regression summary not provided.",
    "reg_coefs_md": (reg_coefs.head(20).to_markdown(index=False) if isinstance(reg_coefs, pd.DataFrame) else "Coefficient table not provided."),
    "methods_text": default_methods(include_fwll_note),
    "limitations_text": default_limitations(),
}

# ---------- Offline template draft ----------
def build_markdown_article(title, authors, affil, date_str, target_venue, length, tone, sections, include_warn: bool, ctx: dict):
    word_hint = {
        "Short (~800 words)": "approximately 800 words",
        "Standard (~1500–2000 words)": "approximately 1500–2000 words",
        "Long (~3000–4000 words)": "approximately 3000–4000 words",
    }.get(length, "approximately 1500–2000 words")

    tone_desc = {
        "Neutral-academic": "neutral, formal academic tone",
        "Didactic/teaching": "didactic, explanatory tone with brief guidance",
        "Critical/replication": "replication-style critical tone emphasizing robustness",
    }.get(tone, "neutral, formal academic tone")

    md = []
    md.append(f"# {title}\n")
    md.append(f"**{authors}**  \n*{affil}*  \n{date_str}\n")
    md.append(f"\n> *Target style:* {target_venue}. *Length:* {word_hint}. *Tone:* {tone_desc}.\n")

    if "Abstract" in sections:
        md.append("## Abstract\n")
        abs_lines = [
            "We present an educational pipeline that juxtaposes exploratory specification search (HARKing) with confirmatory OLS modeling. ",
            "Using transforms, lags, and subset filters, the exploratory step surfaces candidate relationships, while the confirmatory step provides diagnostics and coefficient stability checks. ",
            "Our aim is to illustrate the ease of false discovery under multiple testing and the importance of disciplined validation. ",
        ]
        if include_warn:
            abs_lines.append("We caution that exploratory significance is not evidence of causality and that appropriate multiple-testing controls and theory-driven design remain essential.")
        md.append("".join(abs_lines) + "\n")

    if "Introduction" in sections:
        md.append("## Introduction\n")
        intro = (
            "Empirical analysis in applied economics often begins with uncertain priors and large variable sets. "
            "This context creates fertile ground for HARKing—hypothesizing after results are known—where model selection follows results rather than theory. "
            "To teach these risks constructively, we pair a permissive exploratory grid with a disciplined confirmatory stage. "
            "This article summarizes the outputs of both components and demonstrates how diagnostics inform interpretation."
        )
        md.append(intro + "\n")

    if "Data" in sections:
        md.append("## Data\n")
        md.append(ctx["dataset_summary"] + "\n")

    if "Methods" in sections:
        md.append("## Methods\n")
        md.append(ctx["methods_text"] + "\n")
        if include_warn:
            md.append(
                "\n**Multiple testing and FDR.** Exploratory hits are shown with an FDR flag where available, but the main HARKing table intentionally does **not** adjust p-values to emphasize the false-positive hazard. "
                "Subsequent confirmatory models should prespecify predictors and apply corrections where appropriate.\n"
            )

    if "Results" in sections:
        md.append("## Results\n")
        md.append("### Exploratory HARKing (Top-K snapshot)\n")
        md.append(ctx["hark_summary_md"] + "\n\n")
        md.append("### Multivariate HARKing (Main X + Controls; Top-K snapshot)\n")
        md.append(ctx["multix_summary_md"] + "\n\n")
        md.append("### Confirmatory OLS (summary excerpt)\n")
        md.append("```\n" + textwrap.shorten(ctx["reg_summary_text"], width=5000, placeholder=" …") + "\n```\n\n")
        md.append("**Coefficient table (excerpt)**\n\n")
        md.append(ctx["reg_coefs_md"] + "\n")

    if "Robustness & Diagnostics" in sections:
        md.append("## Robustness & Diagnostics\n")
        rob = (
            "We reviewed residual plots (linearity, homoskedasticity), normal quantile plots, and leverage-residuals maps. "
            "We also computed variance inflation factors (VIF) to screen multicollinearity. Specifications showing systematic residual patterns, "
            "inflated VIF, or leverage-driven fit should be interpreted with caution. Where feasible, we recommend re-estimation on held-out folds or in alternative samples."
        )
        md.append(rob + "\n")

    if "Limitations" in sections:
        md.append("## Limitations\n")
        md.append(ctx["limitations_text"] + "\n")

    if "Conclusion" in sections:
        md.append("## Conclusion\n")
        concl = (
            "Exploration is invaluable for hypothesis generation but must be tempered by validation. "
            "Pairing a transparent search with principled OLS diagnostics clarifies which associations warrant closer scrutiny. "
            "Researchers should pre-register confirmatory specifications, report multiple-testing adjustments, and favor identification strategies when pursuing causal interpretations."
        )
        md.append(concl + "\n")

    if "References" in sections:
        md.append("## References\n")
        md.append(
            "- Gelman, A., & Loken, E. (2014). The garden of forking paths.  \n"
            "- Simmons, J. P., Nelson, L. D., & Simonsohn, U. (2011). False-positive psychology.  \n"
            "- Angrist, J. D., & Pischke, J.-S. (2009). Mostly Harmless Econometrics.  \n"
        )

    return "\n".join(md)

# ---------- Build + Render ----------
st.markdown("#### 3) Build draft")
leftb, rightb = st.columns([1,1])

enable_local_llm = leftb.checkbox(
    "Use Offline LLM (llama.cpp) if model file is present (no API)",
    value=local_llm_available(),
    help="Looks for a GGUF under ./models. Falls back to template if not found.",
    key="aa_use_local_llm"
)
btn_build = rightb.button("🛠️ Build Article Draft", type="primary", key="aa_build")

# store draft in session so you can re-edit
article_md = st.session_state.get("aa_article_md")

if btn_build:
    base_md = build_markdown_article(
        title=title,
        authors=authors,
        affil=affil,
        date_str=date_str,
        target_venue=target_venue,
        length=length,
        tone=tone,
        sections=sections,
        include_warn=include_harking_warning,
        ctx=context
    )
    if enable_local_llm and local_llm_available():
        base_md = refine_with_local_llm(base_md, target_venue, tone, length)
    st.session_state["aa_article_md"] = base_md
    article_md = base_md

st.markdown("---")
st.markdown("#### 4) Draft (editable)")
if article_md:
    article_md = st.text_area("Markdown draft (you can edit before download)", value=article_md, height=500, key="aa_md_editor")
else:
    st.info("Click **Build Article Draft** to generate text.")

# ---------- Downloads ----------
st.markdown("#### 5) Export")
cD1, cD2 = st.columns([1,1])
if article_md:
    # Markdown
    md_bytes = article_md.encode("utf-8")
    cD1.download_button("⬇️ Download Markdown", data=md_bytes, file_name="article_draft.md", mime="text/markdown", key="aa_dl_md")

    # Optional DOCX
    try:
        from docx import Document

        def md_to_docx(md_text: str) -> bytes:
            doc = Document()
            for line in md_text.splitlines():
                if line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                else:
                    doc.add_paragraph(line)
            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue()

        docx_bytes = md_to_docx(article_md)
        cD2.download_button(
            "⬇️ Download DOCX (beta)",
            data=docx_bytes,
            file_name="article_draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="aa_dl_docx"
        )
    except Exception as e:
        cD2.caption(f"DOCX export unavailable ({e}). Install `python-docx` to enable.")

st.markdown("---")
st.caption(
    "Tip: To enable offline LLM drafting, add a small GGUF model under `./models/` and include "
    "`llama-cpp-python` in requirements. If no model is found, the page uses the built-in template only."
)
